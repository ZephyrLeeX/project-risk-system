#!/usr/bin/env python3
"""Build the formal project-risk specification DOCX from the maintained Markdown source."""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "1976D2"
NAVY = "173F5F"
TEXT = "2F4050"
MUTED = "6F879A"
RULE = "D7E3EC"
LIGHT_BLUE = "EAF4FC"
HEADER_FILL = "EDF4F8"
ALT_FILL = "F8FBFD"
RED = "D64545"
PAGE_DXA = 12240
CONTENT_DXA = 9360


def set_run_font(run, size: float, *, bold: bool = False, color: str = TEXT,
                 latin: str = "Calibri", east_asia: str = "Hiragino Sans GB") -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def set_style_font(style, size: float, *, bold: bool = False, color: str = TEXT,
                   latin: str = "Calibri", east_asia: str = "Hiragino Sans GB") -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def choose_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    count = len(headers)
    if count == 2:
        first = 2700 if any(len(r[0]) > 10 for r in rows if r) else 2200
        return [first, CONTENT_DXA - first]
    if count == 3:
        labels = "".join(headers)
        if "默认数据范围" in labels or "主要能力" in labels:
            return [1900, 2200, 5260]
        if "来源" in labels and "说明" in labels:
            return [1900, 2300, 5160]
        return [1800, 2500, 5060]
    if count == 4:
        return [1250, 2050, 3850, 2210]
    if count == 5:
        return [1050, 1650, 2100, 2850, 1710]
    if count == 6:
        return [900, 1450, 1600, 2250, 1700, 1460]
    base = CONTENT_DXA // max(count, 1)
    widths = [base] * count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def set_paragraph_border(paragraph, *, bottom_color=BLUE, bottom_size=18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(bottom_size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), bottom_color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_rpr = OxmlElement("w:rPr")
    field_color = OxmlElement("w:color")
    field_color.set(qn("w:val"), MUTED)
    field_rpr.append(field_color)
    field_run.append(field_rpr)
    for node in (fld_char_begin, instr_text, fld_char_separate, value, fld_char_end):
        field_run.append(node)
    paragraph._p.append(field_run)
    run = paragraph.add_run(" 页")
    set_run_font(run, 9, color=MUTED)


def create_numbering_id(doc: Document, *, bullet: bool = False) -> int:
    """Create a real single-level Word numbering definition that restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "500")
    ind.set(qn("w:hanging"), "260")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    heading_specs = {
        "Title": (24, True, NAVY, 0, 10),
        "Subtitle": (12, False, MUTED, 0, 10),
        "Heading 1": (16, True, BLUE, 16, 8),
        "Heading 2": (13, True, NAVY, 12, 6),
        "Heading 3": (11.5, True, NAVY, 9, 4),
    }
    for name, (size, bold, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, size, bold=bold, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        set_style_font(style, 10.5)
        style.paragraph_format.left_indent = Inches(0.34)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("项目风险管理平台  |  详细规格说明书")
    set_run_font(run, 8.5, bold=True, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    add_page_field(p)

    # Do not display a running header on the title page.
    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目风险管理平台 · V1.0 · 内部评审稿")
    set_run_font(run, 8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    run = p.add_run("PRODUCT SPECIFICATION")
    set_run_font(run, 10, bold=True, color=BLUE)
    run.font.character_spacing = Pt(1.1)

    p = doc.add_paragraph(style="Title")
    p.add_run("项目风险管理平台")
    p = doc.add_paragraph(style="Title")
    run = p.add_run("详细规格说明书")
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph(style="Subtitle")
    run = p.add_run("基于最新 11 个 HTML 原型页面、弹框和跨页交互整理")
    set_run_font(run, 12, color=MUTED)

    doc.add_paragraph()
    metadata = [
        ("文档版本", "V1.0"),
        ("规格基线", "ui-prototype / 01—11 页最新页面"),
        ("编制日期", "2026年7月27日"),
        ("适用阶段", "原型验收、前后端开发、接口联调、测试验收"),
        ("文档状态", "开发基线 / 内部评审稿"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.02)
        run = p.add_run(f"{label}  ")
        set_run_font(run, 9.5, bold=True, color=MUTED)
        run = p.add_run(value)
        set_run_font(run, 11, color=NAVY)
    set_paragraph_border(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    run = p.add_run(
        "本说明书以当前原型为功能和交互基线。原型数据均为演示数据；"
        "正式系统必须接入真实鉴权、数据持久化、邮件、AI、权限和审计服务。"
    )
    set_run_font(run, 10.5, color=TEXT)


def add_manual_toc(doc: Document, headings: list[str]) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(heading)
        set_run_font(run, 10, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("说明：目录列出一级章节；页码以最终渲染版本为准。")
    set_run_font(run, 9, color=MUTED)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str, *, size: float = 10.5, color: str = TEXT) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, max(size - 0.5, 9), color=NAVY, latin="Menlo")
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color=color)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    def split_row(line: str) -> list[str]:
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    headers = split_row(lines[start])
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        row = split_row(lines[index])
        if len(row) < len(headers):
            row += [""] * (len(headers) - len(row))
        rows.append(row[:len(headers)])
        index += 1
    return headers, rows, index


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = choose_widths(headers, rows)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, text, size=9.5, color=NAVY)
        for run in p.runs:
            run.bold = True

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, text in enumerate(values):
            cell = cells[idx]
            if row_index % 2 == 1:
                set_cell_shading(cell, ALT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_inline_runs(p, text, size=9.25)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)


def add_requirement_paragraph(doc: Document, text: str, *, quote: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if quote:
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.18
        set_paragraph_border(p, bottom_color=RULE, bottom_size=5)
        add_inline_runs(p, text, color=NAVY)
    else:
        add_inline_runs(p, text)


def build_doc(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    major_headings = [
        line[3:].strip()
        for line in lines
        if line.startswith("## ") and not line.startswith("### ")
    ]
    add_manual_toc(doc, major_headings)

    index = 0
    active_numbering_id: int | None = None
    active_bullet_id: int | None = None
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if index == 0 and line.startswith("# "):
            index += 1
            continue
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("#### "):
            active_numbering_id = None
            active_bullet_id = None
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            index += 1
            continue
        if line.startswith("### "):
            active_numbering_id = None
            active_bullet_id = None
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("## "):
            active_numbering_id = None
            active_bullet_id = None
            p = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            if line.startswith("## 1."):
                p.paragraph_format.space_before = Pt(0)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\s*\|?\s*:?-{3,}", lines[index + 1]
        ):
            active_numbering_id = None
            active_bullet_id = None
            headers, rows, index = parse_markdown_table(lines, index)
            add_table(doc, headers, rows)
            continue
        match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if match:
            if active_numbering_id is None:
                active_numbering_id = create_numbering_id(doc)
            active_bullet_id = None
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            apply_numbering(p, active_numbering_id)
            add_inline_runs(p, match.group(2))
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            if active_bullet_id is None:
                active_bullet_id = create_numbering_id(doc, bullet=True)
            active_numbering_id = None
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            apply_numbering(p, active_bullet_id)
            add_inline_runs(p, re.sub(r"^[-*]\s+", "", line))
            index += 1
            continue
        if line.startswith("> "):
            active_numbering_id = None
            active_bullet_id = None
            add_requirement_paragraph(doc, line[2:], quote=True)
            index += 1
            continue
        active_numbering_id = None
        active_bullet_id = None
        add_requirement_paragraph(doc, line)
        index += 1

    core = doc.core_properties
    core.title = "项目风险管理平台详细规格说明书"
    core.subject = "基于最新11个HTML原型页面的功能、交互、数据、安全和验收规格"
    core.author = "项目风险管理平台项目组"
    core.keywords = "项目风险, 规格说明书, 原型, 权限, Excel导入, 邮箱同步, API Key"
    core.comments = "以 ui-prototype 01—11 页为当前规格基线"
    core.created = datetime(2026, 7, 27, 0, 0, 0)
    core.modified = datetime(2026, 7, 27, 0, 0, 0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_doc(args.markdown.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
